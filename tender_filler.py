#!/usr/bin/env python3
"""
Tender Form Filler CLI — Fills Italian procurement forms automatically.
Supports .docx (python-docx) and .pdf (pymupdf) formats.

Usage:
    python tender_filler.py --form "path/to/form.docx" --profile company_profile.json
    python tender_filler.py --form "path/to/form.pdf" --profile company_profile.json

Output: Creates a filled copy with "_COMPILATO" suffix in the same directory.

Requirements:
    pip install python-docx pymupdf lxml
"""

import argparse
import json
import os
import re
import sys
import copy
from pathlib import Path
import csv

# ─────────────────────────────────────────────
# DEBUG CONFIGURATION
# ─────────────────────────────────────────────
VERBOSE_MODE = False  # Set to True for detailed debug logging

# ─────────────────────────────────────────────
# SEMANTIC FIELD MAPPING
# Maps Italian form labels to company profile keys
# ─────────────────────────────────────────────

SEMANTIC_MAP = [
    # Person / Legal Representative
    (r"(?:il |la )?sottoscritt[oa]|nome\s*(?:e\s*)?cognome|cognome\s*e\s*nome|legale\s*rappresentante|rappresentante\s*legale|titolare|(?:sig|dott)\.\s*\/?\s*(?:sig|dott)\.?ra",
     "legale_rappresentante.nome_completo"),
    (r"nat[oa]\s+a|luogo\s*(?:di\s*)?nascita|comune\s*(?:di\s*)?nascita",
     "legale_rappresentante.luogo_nascita"),
    (r"(?:il|in\s*data)\s*\d|data\s*(?:di\s*)?nascita|nato\s*il",
     "legale_rappresentante.data_nascita"),
    (r"prov(?:\.|incia)?\s*(?:di\s*)?nascita|prov\.\s*\(",
     "legale_rappresentante.provincia_nascita"),
    (r"c(?:odice|\.?\s*)f(?:iscale|\.?\s*)(?:persona|personale|individuale|del\s*(?:legale|rappresentante|dichiarante|sottoscritto))?",
     "legale_rappresentante.codice_fiscale"),
    (r"residen(?:za|te)\s*(?:in|a)|domicilio|indirizzo\s*residen",
     "legale_rappresentante.residenza"),
    (r"in\s*qualit[aà]\s*di|carica|ruolo|qualifica",
     "legale_rappresentante.qualifica"),

    # Company
    (r"ragione\s*sociale|denominazione\s*(?:sociale|impresa|azienda)|ditta|impresa|societ[aà]|operatore\s*economico",
     "azienda.ragione_sociale"),
    (r"forma\s*giuridica|natura\s*giuridica|tipo\s*societ",
     "azienda.forma_giuridica"),
    (r"c(?:odice|\.?\s*)f(?:iscale|\.?\s*)(?:aziend|societ|impresa|ditta)|codice\s*fiscale\s*$|c\.?f\.?\s*$",
     "azienda.cf_piva"),
    (r"p(?:artita|\.?\s*)i(?:va|\.?\s*v\.?\s*a\.?\s*)|partita\s*iva",
     "azienda.cf_piva"),
    (r"sede\s*legale(?!\s*amm)|sede\s*(?:in|a)\s*\(|con\s*sede\s*(?:legale|in)|via\s*dell",
     "azienda.sede_legale"),
    (r"cap\s*(?:sede\s*legale)?|c\.?a\.?p\.?\s*(?:legale|$)",
     "azienda.sede_legale_cap"),
    (r"citt[aà]\s*sede\s*legale|comune\s*sede\s*legale|localit[aà]",
     "azienda.sede_legale_citta"),
    (r"prov(?:\.|incia)?\s*(?:di\s*)?sede\s*legale|prov\.\s*sede",
     "azienda.sede_legale_provincia"),
    (r"sede\s*amministrativa|sede\s*amm(?:\.|istrativa)?",
     "azienda.sede_amministrativa"),
    (r"cap\s*(?:amministrativa|amm\.?)|c\.?a\.?p\.?\s*amm",
     "azienda.sede_amm_cap"),
    (r"citt[aà]\s*(?:amministrativa|amm\.?)|comune\s*(?:amministrativa|amm\.?)",
     "azienda.sede_amm_citta"),
    (r"prov(?:\.|incia)?\s*(?:amministrativa|amm\.?)",
     "azienda.sede_amm_provincia"),
    (r"tel(?:efono|\.?\s*)|phone|recapito\s*telefonico",
     "azienda.telefono"),
    (r"fax",
     "azienda.fax"),
    (r"p\.?e\.?c\.?\s*(?::|$)|posta\s*elettronica\s*certificata",
     "azienda.pec"),
    (r"e-?mail(?:\s*ordinaria)?|posta\s*elettronica\s*(?!cert)",
     "azienda.email"),

    # Registry
    (r"c\.?c\.?i\.?a\.?a\.?|camera\s*(?:di\s*)?commercio",
     "azienda.cciaa"),
    (r"r\.?e\.?a\.?\s*(?:n\.?|numero)?",
     "azienda.rea"),
    (r"data\s*(?:di\s*)?iscrizione|iscritt[oa]\s*(?:il|in\s*data|dal)",
     "azienda.data_iscrizione"),
    (r"capitale\s*sociale",
     "azienda.capitale_sociale"),
    (r"ateco|codice\s*attivit[aà]",
     "azienda.ateco"),
    (r"c\.?c\.?n\.?l\.?|contratto\s*(?:collettivo|nazionale)",
     "azienda.ccnl"),
    (r"(?:n(?:umero|\.?)?\s*)?dipendenti|organico|personale\s*(?:medio|complessivo)|occupati|numero\s+complessivo|complessivi",
     "azienda.dipendenti_totale"),
    (r"(?:attivit[aà]|oggetto|descrizione|settore|ramo|classe|natura)\s*(?:economica|principale|dell[a'])?|descrizione\s*(?:attivit[aà]|della\s*(?:attivit[aà]|impresa))",
     "azienda.ateco_descrizione"),
    
    # Additional patterns for common variations
    (r"(?:indirizzo|via|corso|piazza|viale|largo|borgo|contrada)\s*(?:dell[a']|del|di|della)?\s*[A-Z]",
     "azienda.sede_legale"),
    (r"n\.?\s*civico|numero\s*civico|snc",
     "azienda.sede_legale"),
    (r"cognome|nome\s*e\s*cognome\s*del\s*sottoscritto",
     "legale_rappresentante.nome_completo"),
    (r"qual(?:ifica|it[aà])\s*(?:del\s*)?(?:legale\s*)?(?:rappresentante|dichiarante)",
     "legale_rappresentante.qualifica"),
    (r"telefono|cellulare|contatto|numero\s*di\s*telefono",
     "azienda.telefono"),
    (r"indirizzo\s*e-?mail|email\s*ordinaria|posta\s*ordinaria",
     "azienda.email"),
    
    # Explicit full forms (highest priority)
    (r"ragione\s*sociale\s*dell[a'](?:impresa|azienda|ditta|societ[aà])?|denominazione\s*della\s*(?:impresa|azienda|ditta|societ[aà])",
     "azienda.ragione_sociale"),
    (r"numero\s+(?:di\s+)?dipendenti|complessivo\s+di\s+dipendenti|numero\s+complessivo",
     "azienda.dipendenti_totale"),
    (r"natura\s+dell[a']\s*attivit[aà]|descrizione\s+dell[a']\s*(?:attivit[aà]|impresa)|oggetto\s+dell[a']\s*attivit[aà]|attivit[aà]\s+principale",
     "azienda.ateco_descrizione"),
]


def load_profile(path):
    """Load company profile from JSON or CSV file."""
    ext = Path(path).suffix.lower()
    if ext == '.json':
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    elif ext == '.csv':
        profile = {}
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                section = row.get('section', '').strip()
                key = row.get('key', '').strip()
                value = row.get('value', '').strip()
                if not section or not key:
                    continue
                if section not in profile:
                    profile[section] = {}
                # Handle lists like 'soci'
                if section == 'soci':
                    if 'soci' not in profile:
                        profile['soci'] = []
                    profile['soci'].append({
                        'nome': row.get('nome', value),
                        'quota': row.get('quota', ''),
                        'ruolo': row.get('ruolo', '')
                    })
                else:
                    profile[section][key] = value
        return profile
    else:
        raise ValueError(f"Unsupported profile format: {ext}. Use .json or .csv")


def get_profile_value(profile, dotted_key):
    """Retrieve a value from nested profile dict using dotted key notation."""
    keys = dotted_key.split('.')
    val = profile
    for k in keys:
        if isinstance(val, dict) and k in val:
            val = val[k]
        else:
            return None
    return val


def match_label(text, profile):
    """
    Given a text label, find the best matching profile value using semantic map.
    Returns (matched_key, value) or (None, None).
    """
    text_lower = text.lower().strip()
    if len(text_lower) < 2:
        return None, None

    for pattern, key in SEMANTIC_MAP:
        if re.search(pattern, text_lower, re.IGNORECASE):
            val = get_profile_value(profile, key)
            if val:
                if VERBOSE_MODE:
                    print(f"      ✓ Matched '{text[:50]}' → {key} = '{str(val)[:40]}'")
                return key, val
    
    if VERBOSE_MODE and len(text_lower) > 3:
        print(f"      ✗ No match for '{text[:60]}'")
    return None, None


def highlight_empty_fields(doc):
    """
    Scan document for truly unfilled fields and highlight them in yellow.
    Skip fields that have been filled with actual content.
    Returns count of highlighted fields.
    """
    from docx.enum.text import WD_COLOR_INDEX
    count = 0
    
    for para in doc.paragraphs:
        # Scan each run for unfilled markers that are NOT part of filled content
        for run in para.runs:
            if not run.text:
                continue
            
            # Check if this run contains ONLY placeholder markers (no real content)
            text = run.text
            
            # Pattern: pure underscore blanks (3+ underscores, maybe with trailing/leading space)
            if re.match(r'^\s*_{3,}\s*$', text):
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    count += 1
                except:
                    pass
            # Pattern: pure dots/ellipsis (3+ dots, maybe spaces)
            elif re.match(r'^\s*(?:\.{3,}|…+)\s*$', text):
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    count += 1
                except:
                    pass
    
    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text:
                            continue
                        
                        text = run.text
                        
                        # Only highlight pure placeholder runs
                        if re.match(r'^\s*_{3,}\s*$', text):
                            try:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                count += 1
                            except:
                                pass
                        elif re.match(r'^\s*(?:\.{3,}|…+)\s*$', text):
                            try:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                count += 1
                            except:
                                pass
    
    return count


# ═════════════════════════════════════════════
# DOCX FILLING ENGINE
# ═════════════════════════════════════════════

def fill_docx(form_path, profile, output_path):
    """Fill a .docx form using multiple strategies."""
    try:
        from docx import Document
        from lxml import etree
    except ImportError:
        print("ERROR: Missing dependencies. Run: pip install python-docx lxml")
        sys.exit(1)

    doc = Document(form_path)
    stats = {"underscore": 0, "formtext": 0, "sdt_checkbox": 0, "table_cell": 0, "context_run": 0}

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    }
    W = ns['w']
    W14 = ns['w14']

    # Short-label map: tiny labels that appear right before a blank
    # These are resolved by position relative to other filled fields
    SHORT_LABEL_MAP = {
        r'\bprov\.?\b': 'province_of_previous',   # province follows a city/birthplace
        r'\bil\b\s*$': 'date_of_previous',         # "il" = date follows a birthplace
        r'\bcitt[aà]\b': 'azienda.sede_legale_citta',
        r'\bcap\b': 'azienda.sede_legale_cap',
        r'\bcomune\b': 'azienda.sede_legale_citta',
    }

    # ── Strategy 1: Underscore blank replacement ──
    def fill_underscore_blanks(paragraph):
        """Join all runs, find each underscore blank, use the ORIGINAL preceding
        label text (not the filled value) for matching."""
        full_text = "".join(r.text for r in paragraph.runs if r.text)
        if not full_text or '___' not in full_text:
            return False

        # Split into segments: text and underscore blocks alternate
        # Even indices = text, odd indices = underscore blocks
        segments = re.split(r'(_{3,})', full_text)
        result_segments = list(segments)  # Copy — we modify this one
        filled = False
        used_keys = set()
        last_matched_key = None

        for i, seg in enumerate(segments):
            if not re.match(r'^_{3,}$', seg):
                continue  # Only process underscore segments

            # Get the ORIGINAL preceding text segment (always at i-1 in original segments)
            preceding_text = segments[i - 1].strip() if i > 0 else ""

            # First: try the standard semantic map
            key, val = match_label(preceding_text, profile)
            if key and key in used_keys:
                key, val = None, None

            # Second: try short-label contextual resolution
            if not val:
                preceding_lower = preceding_text.lower().strip()
                for short_pat, short_action in SHORT_LABEL_MAP.items():
                    if re.search(short_pat, preceding_lower, re.IGNORECASE):
                        if short_action == 'province_of_previous':
                            if last_matched_key and 'nascita' in last_matched_key:
                                val = get_profile_value(profile, 'legale_rappresentante.provincia_nascita')
                                key = 'legale_rappresentante.provincia_nascita'
                            elif last_matched_key and 'sede' in last_matched_key:
                                val = get_profile_value(profile, 'azienda.sede_legale_provincia')
                                key = 'azienda.sede_legale_provincia'
                            else:
                                val = get_profile_value(profile, 'legale_rappresentante.provincia_nascita')
                                key = 'legale_rappresentante.provincia_nascita'
                        elif short_action == 'date_of_previous':
                            if last_matched_key and 'nascita' in last_matched_key:
                                val = get_profile_value(profile, 'legale_rappresentante.data_nascita')
                                key = 'legale_rappresentante.data_nascita'
                            elif last_matched_key and 'iscrizi' in last_matched_key:
                                val = get_profile_value(profile, 'azienda.data_iscrizione')
                                key = 'azienda.data_iscrizione'
                        else:
                            val = get_profile_value(profile, short_action)
                            key = short_action
                        if val and key not in used_keys:
                            break
                        else:
                            key, val = None, None

            if val and key not in used_keys:
                result_segments[i] = str(val)
                used_keys.add(key)
                last_matched_key = key
                filled = True

        if filled:
            new_text = "".join(result_segments)
            # Fix duplicates like "Via Via dell'Industria"
            for section in profile.values():
                if isinstance(section, dict):
                    for v in section.values():
                        v_str = str(v)
                        words = v_str.split()
                        if words:
                            dup = f"{words[0]} {v_str}"
                            if dup in new_text:
                                new_text = new_text.replace(dup, v_str)

            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for r in paragraph.runs[1:]:
                    r.text = ""
                stats["underscore"] += 1
            return True
        return False

    # ── Strategy 2: FORMTEXT field filling ──
    def fill_formtext_in_element(element, context_text=""):
        """Find and fill FORMTEXT (fldChar) fields."""
        runs = element.findall(f'.//{{{W}}}r')
        in_formtext = False
        found_separate = False
        current_context = context_text

        for run in runs:
            # Check for field char
            fc = run.find(f'{{{W}}}fldChar')
            if fc is not None:
                ftype = fc.get(f'{{{W}}}fldCharType')
                if ftype == 'begin':
                    in_formtext = False
                    found_separate = False
                elif ftype == 'separate' and in_formtext:
                    found_separate = True
                    continue
                elif ftype == 'end':
                    in_formtext = False
                    found_separate = False
                    continue

            # Check for instrText
            instr = run.find(f'{{{W}}}instrText')
            if instr is not None and instr.text and 'FORMTEXT' in instr.text:
                in_formtext = True
                continue

            # Fill the value run after 'separate'
            if found_separate:
                t_el = run.find(f'{{{W}}}t')
                if t_el is not None:
                    key, val = match_label(current_context, profile)
                    if val:
                        t_el.text = str(val)
                        stats["formtext"] += 1
                    found_separate = False

            # Accumulate context from regular text
            t_el = run.find(f'{{{W}}}t')
            if t_el is not None and t_el.text and not in_formtext:
                current_context += " " + t_el.text

    # ── Strategy 3: SDT checkbox toggling ──
    def process_sdt_checkboxes(element):
        """Find SDT checkbox elements and check relevant ones."""
        sdts = element.findall(f'.//{{{W}}}sdt')
        for sdt in sdts:
            # Is this a checkbox?
            checkbox = sdt.find(f'.//{{{W14}}}checkbox')
            if checkbox is None:
                continue

            # Get surrounding context to decide if it should be checked
            # Look for text in the same paragraph
            parent = sdt.getparent()
            if parent is not None:
                context_texts = []
                for t in parent.iter(f'{{{W}}}t'):
                    if t.text:
                        context_texts.append(t.text)
                context = " ".join(context_texts).lower()

                # Decide what to check based on context
                should_check = False

                # Common patterns in Italian tenders
                if any(kw in context for kw in ['singol', 'individuale', 'impresa singola']):
                    should_check = True  # Typically a single company
                elif 'subappalto' in context:
                    should_check = False  # Default: no subcontracting
                elif 'avvalimento' in context:
                    should_check = False  # Default: no avvalimento
                elif any(kw in context for kw in ['consorziat', 'raggruppamento', 'rti', 'ati']):
                    should_check = False  # Default: not a consortium
                elif any(kw in context for kw in ['microimpresa', 'piccola', 'media']):
                    # Check based on company size
                    dip = profile.get('azienda', {}).get('dipendenti_totale', '0')
                    try:
                        n = int(dip)
                        if 'micro' in context and n < 10:
                            should_check = True
                        elif 'piccola' in context and 10 <= n < 50:
                            should_check = True
                        elif 'media' in context and 50 <= n < 250:
                            should_check = True
                    except ValueError:
                        pass

                if should_check:
                    checked_el = checkbox.find(f'{{{W14}}}checked')
                    if checked_el is not None:
                        checked_el.set(f'{{{W14}}}val', '1')
                    else:
                        checked_el = etree.SubElement(checkbox, f'{{{W14}}}checked')
                        checked_el.set(f'{{{W14}}}val', '1')

                    # Update display character
                    for t in sdt.iter(f'{{{W}}}t'):
                        t.text = '\u2612'  # ☒

                    stats["sdt_checkbox"] += 1

    # ── Strategy 4: Table cell label-value filling ──
    def fill_tables(doc):
        """Fill table cells where label is in one cell and value in another."""
        for table in doc.tables:
            for row in table.rows:
                # De-duplicate merged cells
                seen_ids = set()
                unique_cells = []
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid not in seen_ids:
                        seen_ids.add(cid)
                        unique_cells.append(cell)

                if len(unique_cells) < 2:
                    continue

                # Check pairs: label cell → value cell
                for i in range(len(unique_cells) - 1):
                    label_text = unique_cells[i].text.strip()
                    value_cell = unique_cells[i + 1]
                    value_text = value_cell.text.strip()

                    key, val = match_label(label_text, profile)
                    if val:
                        # Check if value cell has FORMTEXT
                        has_formtext = False
                        for p in value_cell.paragraphs:
                            for run in p._element.findall(f'.//{{{W}}}r'):
                                instr = run.find(f'{{{W}}}instrText')
                                if instr is not None and instr.text and 'FORMTEXT' in instr.text:
                                    has_formtext = True
                                    break

                        if has_formtext:
                            fill_formtext_in_element(value_cell._tc, label_text)
                        else:
                            # Check if it has example/placeholder data
                            is_placeholder = (
                                not value_text or
                                value_text in ['...', '…', '____', '________'] or
                                value_text == label_text or
                                len(value_text) < 3
                            )
                            # Also replace known example data
                            example_names = ['PIETRO FIORENTINI', 'MARIO PIETRO NARDI',
                                           'NOME COGNOME', 'DENOMINAZIONE']
                            is_example = any(ex in value_text.upper() for ex in example_names)

                            if is_placeholder or is_example:
                                for p in value_cell.paragraphs:
                                    # Preserve formatting of first run
                                    if p.runs:
                                        p.runs[0].text = str(val)
                                        for r in p.runs[1:]:
                                            r.text = ""
                                    else:
                                        p.text = str(val)
                                    stats["table_cell"] += 1
                                    break  # Only fill first paragraph

    # ── Strategy 5: Context-aware run replacement ──
    def fill_context_runs(paragraph):
        """For paragraphs with inline labels followed by values on the same line."""
        full_text = "".join(r.text for r in paragraph.runs if r.text)
        if not full_text:
            return

        # Pattern: "Label: ..." or "Label: VALUE" where VALUE is placeholder
        patterns = [
            (r'((?:ragione\s*sociale|denominazione)[:\s]*)\s*\.{2,}', 'azienda.ragione_sociale'),
            (r'((?:c\.?\s*f\.?|codice\s*fiscale)[:\s]*)\s*\.{2,}', 'azienda.cf_piva'),
            (r'((?:p\.?\s*i\.?\s*v\.?\s*a\.?|partita\s*iva)[:\s]*)\s*\.{2,}', 'azienda.cf_piva'),
            (r'((?:sede\s*legale|sede\s*in|con\s*sede)[:\s]*)\s*\.{2,}', 'azienda.sede_legale'),
            (r'((?:pec)[:\s]*)\s*\.{2,}', 'azienda.pec'),
            (r'((?:tel(?:efono)?\.?)[:\s]*)\s*\.{2,}', 'azienda.telefono'),
            (r'((?:fax)[:\s]*)\s*\.{2,}', 'azienda.fax'),
        ]

        modified = full_text
        changed = False
        for pat, key in patterns:
            val = get_profile_value(profile, key)
            if val:
                new_text = re.sub(pat, rf'\1 {val}', modified, flags=re.IGNORECASE, count=1)
                if new_text != modified:
                    modified = new_text
                    changed = True

        if changed:
            # Clean up any trailing underscores or dots after the filled value
            modified = re.sub(r'(\S+(?:\s+\S+)*)\s*_{3,}', r'\1', modified)  # Remove trailing underscores
            modified = re.sub(r'(\S+(?:\s+\S+)*)\s*\.{3,}', r'\1', modified)  # Remove trailing dots
            
            if paragraph.runs:
                paragraph.runs[0].text = modified
                for r in paragraph.runs[1:]:
                    r.text = ""
                stats["context_run"] += 1

    # ── Execute all strategies ──
    print("  [DOCX] Running Strategy 1: Underscore blank replacement...")
    for para in doc.paragraphs:
        fill_underscore_blanks(para)

    print("  [DOCX] Running Strategy 2: FORMTEXT field filling...")
    for para in doc.paragraphs:
        fill_formtext_in_element(para._element, para.text)

    print("  [DOCX] Running Strategy 3: SDT checkbox toggling...")
    process_sdt_checkboxes(doc.element)

    print("  [DOCX] Running Strategy 4: Table cell filling...")
    fill_tables(doc)

    # Also check tables for FORMTEXT and checkboxes
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fill_formtext_in_element(para._element, cell.text)
                process_sdt_checkboxes(cell._tc)

    print("  [DOCX] Running Strategy 5: Context-aware run replacement...")
    for para in doc.paragraphs:
        fill_context_runs(para)

    print("  [DOCX] Running Strategy 6: Highlighting unfilled fields in yellow...")
    highlighted_count = highlight_empty_fields(doc)

    # Save
    doc.save(output_path)

    print(f"\n  ✅ DOCX filled successfully!")
    print(f"     Underscore blanks filled: {stats['underscore']}")
    print(f"     FORMTEXT fields filled:   {stats['formtext']}")
    print(f"     SDT checkboxes toggled:   {stats['sdt_checkbox']}")
    print(f"     Table cells filled:       {stats['table_cell']}")
    print(f"     Context runs replaced:    {stats['context_run']}")
    print(f"     Empty fields highlighted: {highlighted_count}")
    return True


# ═════════════════════════════════════════════
# PDF FILLING ENGINE
# ═════════════════════════════════════════════

def fill_pdf(form_path, profile, output_path):
    """Fill a PDF form — handles both interactive (AcroForm) and flat PDFs."""
    try:
        import fitz  # pymupdf
    except ImportError:
        print("ERROR: Missing pymupdf. Run: pip install pymupdf")
        sys.exit(1)

    doc = fitz.open(form_path)
    stats = {"acroform": 0, "overlay": 0, "checkbox": 0}

    # ── Check for interactive form fields first ──
    has_widgets = False
    for page in doc:
        widgets = list(page.widgets())
        if widgets:
            has_widgets = True
            for widget in widgets:
                field_name = widget.field_name or ""
                field_label = widget.field_label or field_name

                # Try to match the field
                key, val = match_label(field_label, profile)
                if not val:
                    key, val = match_label(field_name, profile)

                if val:
                    if widget.field_type == fitz.PDF_WIDGET_TYPE_TEXT:
                        widget.field_value = str(val)
                        widget.update()
                        stats["acroform"] += 1
                    elif widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                        widget.field_value = True
                        widget.update()
                        stats["acroform"] += 1

    if has_widgets:
        doc.save(output_path)
        doc.close()
        print(f"\n  ✅ PDF (interactive form) filled successfully!")
        print(f"     AcroForm fields filled: {stats['acroform']}")
        return True

    # ── Flat PDF: coordinate-based text overlay ──
    print("  [PDF] No interactive fields found — using coordinate overlay...")

    # Build a flat lookup of all profile values
    flat_values = {}
    for section_key, section in profile.items():
        if isinstance(section, dict):
            for k, v in section.items():
                flat_values[f"{section_key}.{k}"] = str(v)

    for page_num, page in enumerate(doc):
        page_text = page.get_text("text")
        blocks = page.get_text("dict")["blocks"]

        # Build list of text positions for context
        text_positions = []
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text_positions.append({
                            "text": span["text"],
                            "bbox": span["bbox"],  # (x0, y0, x1, y1)
                            "size": span["size"],
                            "font": span["font"],
                        })

        # Find labels and overlay values after them
        for i, tp in enumerate(text_positions):
            label = tp["text"].strip()
            if len(label) < 3:
                continue

            key, val = match_label(label, profile)
            if not val:
                continue

            # Check if there's already a value after this label
            bbox = tp["bbox"]
            x_after = bbox[2] + 5  # 5px after label ends
            y_pos = bbox[3]  # baseline

            # Look if the next span on the same line already has the value
            already_filled = False
            if i + 1 < len(text_positions):
                next_tp = text_positions[i + 1]
                next_y = next_tp["bbox"][1]
                # Same line (within 3px tolerance)
                if abs(next_y - bbox[1]) < 3:
                    next_text = next_tp["text"].strip()
                    # Skip if already has real content (not just dots/underscores/spaces)
                    if next_text and not re.match(r'^[._\s\-:]+$', next_text):
                        # Check if it's the same value we'd fill — already done
                        if next_text == str(val):
                            already_filled = True
                        # Otherwise it might be part of the label — fill after it
                        else:
                            x_after = next_tp["bbox"][2] + 5

            if not already_filled:
                # Cover any dots/underscores with white rectangle
                line_end_x = page.rect.width - 50  # Leave margin
                cover_rect = fitz.Rect(x_after - 2, bbox[1] - 1, line_end_x, bbox[3] + 1)
                page.draw_rect(cover_rect, color=None, fill=(1, 1, 1))

                # Overlay the value
                fontsize = tp["size"] if tp["size"] > 6 else 10
                page.insert_text(
                    fitz.Point(x_after, y_pos),
                    str(val),
                    fontname="helv",
                    fontsize=fontsize,
                    color=(0, 0, 0)
                )
                stats["overlay"] += 1
                print(f"     Page {page_num+1}: '{label[:40]}' → '{val}'")

    doc.save(output_path)
    doc.close()

    print(f"\n  ✅ PDF (flat overlay) filled successfully!")
    print(f"     Text overlays placed: {stats['overlay']}")
    return True


# ═════════════════════════════════════════════
# CLI ENTRY POINT
# ═════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Tender Form Filler — Automatically fills Italian procurement forms",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python tender_filler.py --form "Allegato_1.docx" --profile company_profile.json
  python tender_filler.py --form "Domanda.pdf" --profile company_profile.csv
  python tender_filler.py --auto  (processes all forms in EMPTY_FORM/ and saves to FILLED_FORM/)
        """
    )
    parser.add_argument('--form', '-f', help='Path to the form file (.docx or .pdf)')
    parser.add_argument('--profile', '-p', default='company_profile.json',
                       help='Path to company profile JSON/CSV (default: company_profile.json)')
    parser.add_argument('--output', '-o', default=None,
                       help='Output file path (default: adds _COMPILATO suffix)')
    parser.add_argument('--auto', action='store_true',
                       help='Automatically process all forms in EMPTY_FORM/ and save to FILLED_FORM/')
    parser.add_argument('--verbose', '-v', action='store_true',
                       help='Enable detailed debug logging for field matching')

    args = parser.parse_args()
    
    # Set global verbose mode
    global VERBOSE_MODE
    VERBOSE_MODE = args.verbose

    # Load profile
    profile_path = Path(args.profile)
    if not profile_path.exists():
        # Try looking in script directory
        script_dir = Path(__file__).parent
        profile_path = script_dir / args.profile
        if not profile_path.exists():
            print(f"❌ Profile file not found: {args.profile}")
            print(f"   Create one by copying and editing company_profile.json or .csv")
            sys.exit(1)

    print(f"\n📋 Loading company profile from: {profile_path}")
    try:
        profile = load_profile(profile_path)
    except Exception as e:
        print(f"❌ Error loading profile: {e}")
        sys.exit(1)

    company = profile.get('azienda', {}).get('ragione_sociale', 'Unknown')
    rep = profile.get('legale_rappresentante', {}).get('nome_completo', 'Unknown')
    print(f"   Company: {company}")
    print(f"   Representative: {rep}")

    if args.auto:
        # Auto mode: process all forms in EMPTY_FORM/
        empty_dir = Path('EMPTY_FORM')
        filled_dir = Path('FILLED_FORM')
        filled_dir.mkdir(exist_ok=True)

        if not empty_dir.exists():
            print("❌ EMPTY_FORM/ directory not found")
            sys.exit(1)

        forms = list(empty_dir.glob('*.docx')) + list(empty_dir.glob('*.pdf'))
        if not forms:
            print("❌ No .docx or .pdf files found in EMPTY_FORM/")
            sys.exit(1)

        print(f"\n📄 Processing {len(forms)} forms from EMPTY_FORM/ to FILLED_FORM/\n")

        for form_path in forms:
            output_path = filled_dir / f"{form_path.stem}_COMPILATO{form_path.suffix}"
            print(f"Processing: {form_path.name} → {output_path.name}")
            process_form(form_path, profile, output_path)

        print(f"\n🎉 All forms processed! Check FILLED_FORM/ for results.")
    else:
        # Manual mode: single form
        if not args.form:
            print("❌ Specify --form or use --auto")
            sys.exit(1)

        form_path = Path(args.form)
        if not form_path.exists():
            print(f"❌ Form file not found: {form_path}")
            sys.exit(1)

        # Determine output path
        if args.output:
            output_path = Path(args.output)
        else:
            suffix = form_path.suffix
            output_path = form_path.with_name(f"{form_path.stem}_COMPILATO{suffix}")

        print(f"\n📄 Processing: {form_path.name}")
        print(f"   Output: {output_path.name}\n")

        process_form(form_path, profile, output_path)

        print(f"\n🎉 Done! Filled form saved to: {output_path}")


def process_form(form_path, profile, output_path):
    """Process a single form file."""
    ext = form_path.suffix.lower()
    if ext == '.docx':
        success = fill_docx(str(form_path), profile, str(output_path))
    elif ext == '.pdf':
        success = fill_pdf(str(form_path), profile, str(output_path))
    else:
        print(f"❌ Unsupported file format: {ext}")
        print(f"   Supported: .docx, .pdf")
        return False

    if not success:
        print(f"⚠️  Form processing completed with warnings. Check output: {output_path}")
    return success


if __name__ == '__main__':
    main()
